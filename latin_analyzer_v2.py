#!/usr/bin/env python3
"""
Analyseur de textes latins médiévaux - Version 2
================================================

Système intégré utilisant :
- PyCollatinus pour l'analyse morphologique
- Dictionnaire Du Cange (99k+ entrées de latin médiéval)
- Système de scoring multi-critères
- Colorisation à 3 niveaux (rouge/orange/noir)

Auteur: Claude
Date: 2025-11-24
"""

import os
import re
import sys
import docx
from docx.shared import RGBColor, Cm, Pt
from docx.enum.text import WD_LINE_SPACING
from pathlib import Path
from collections import Counter

# Ajouter PyCollatinus au path
sys.path.insert(0, '/tmp/collatinus-python')

from pycollatinus import Lemmatiseur
from page_xml_parser import PageXMLParser


class LatinAnalyzer:
    """Analyseur de textes latins avec détection intelligente des erreurs."""

    def __init__(self, ducange_dict_file=None):
        """
        Initialise l'analyseur.

        Args:
            ducange_dict_file (str): Chemin vers le dictionnaire Du Cange
        """
        print("🔄 Initialisation de l'analyseur latin...")

        # Charger PyCollatinus
        print("  📚 Chargement de Collatinus (latin classique)...")
        self.lemmatizer = Lemmatiseur()
        print("  ✅ Collatinus chargé")

        # Charger le dictionnaire médiéval
        self.medieval_dict = set()
        if ducange_dict_file and os.path.exists(ducange_dict_file):
            print(f"  📚 Chargement du dictionnaire Du Cange...")
            with open(ducange_dict_file, 'r', encoding='utf-8') as f:
                for line in f:
                    word = line.strip().lower()
                    if word:
                        self.medieval_dict.add(word)
            print(f"  ✅ {len(self.medieval_dict)} mots médiévaux chargés")
        else:
            print("  ⚠️  Dictionnaire Du Cange non trouvé")

        # Suffixes typiques du latin médiéval
        self.medieval_suffixes = [
            'arius', 'aria', 'arium',  # Noms d'agent
            'atio', 'ationis',  # Noms d'action
            'tor', 'toris',  # Agent
            'torium', 'torii',  # Lieu
            'mentum', 'menti',  # Instrument
            'itia', 'itiae',  # Qualité
        ]

        # Contexte ecclésiastique
        self.ecclesiastical_words = {
            'abbas', 'abbatia', 'abbatissa', 'archiepiscopus', 'basilica',
            'canonicus', 'capitulum', 'cardinalis', 'clericus', 'diaconus',
            'diocesis', 'dominus', 'ecclesia', 'episcopus', 'monasterium',
            'monachus', 'parochia', 'presbyter', 'sacerdos', 'sanctus',
        }

        print("✅ Analyseur prêt\n")

    def analyze_word(self, word, context_words=None):
        """
        Analyse un mot avec scoring multi-critères.

        Args:
            word (str): Le mot à analyser
            context_words (list): Liste des mots environnants

        Returns:
            dict: {
                'word': str,
                'recognized_classical': bool,
                'recognized_medieval': bool,
                'confidence_score': int (0-100),
                'reasons': list,
                'color_code': str ('black', 'orange', 'red')
            }
        """
        clean_word = word.lower().strip()
        context_words = context_words or []

        result = {
            'word': word,
            'recognized_classical': False,
            'recognized_medieval': False,
            'confidence_score': 50,  # Score de base neutre
            'reasons': [],
            'color_code': 'orange'
        }

        # Ignorer les mots très courts
        if len(clean_word) < 2:
            result['confidence_score'] = 100
            result['color_code'] = 'black'
            result['reasons'].append("mot trop court pour être analysé")
            return result

        # Critère 1 : Reconnu par Collatinus (latin classique)
        try:
            analyses = self.lemmatizer.lemmatise(clean_word)
            if analyses and len(analyses) > 0:
                result['recognized_classical'] = True
                result['confidence_score'] += 30
                result['reasons'].append(f"latin classique valide ({len(analyses)} analyse(s))")
        except Exception:
            pass

        # Critère 2 : Présent dans le dictionnaire Du Cange
        if clean_word in self.medieval_dict:
            result['recognized_medieval'] = True
            result['confidence_score'] += 40
            result['reasons'].append("présent dans le dictionnaire Du Cange")

        # Critère 3 : Suffixe médiéval typique
        for suffix in self.medieval_suffixes:
            if clean_word.endswith(suffix):
                result['confidence_score'] += 10
                result['reasons'].append(f"suffixe médiéval productif (-{suffix})")
                break

        # Critère 4 : Contexte ecclésiastique
        context_ecclesiastical = any(
            w.lower() in self.ecclesiastical_words
            for w in context_words
        )
        if context_ecclesiastical:
            result['confidence_score'] += 5
            result['reasons'].append("contexte ecclésiastique")

        # Critère 5 : Variante orthographique médiévale (ae→e, ti→ci, etc.)
        if self._is_medieval_variant(clean_word):
            result['confidence_score'] += 10
            result['reasons'].append("variante orthographique médiévale détectée")

        # Déterminer la couleur selon le score
        score = min(result['confidence_score'], 100)
        if score >= 75:
            result['color_code'] = 'black'
        elif score >= 40:
            result['color_code'] = 'orange'
        else:
            result['color_code'] = 'red'

        result['confidence_score'] = score
        return result

    def _is_medieval_variant(self, word):
        """
        Détecte si un mot pourrait être une variante orthographique médiévale.

        Args:
            word (str): Le mot à analyser

        Returns:
            bool: True si variante détectée
        """
        # Générer des variantes classiques et vérifier si elles sont reconnues
        variants = []

        # ae ← e (medieval → mediaeval)
        if 'e' in word and 'ae' not in word:
            variants.append(word.replace('e', 'ae', 1))

        # ti ← ci (gracia → gratia)
        if 'ci' in word:
            variants.append(word.replace('ci', 'ti'))

        # Vérifier si une variante existe dans les dictionnaires
        for variant in variants:
            if variant in self.medieval_dict:
                return True
            try:
                analyses = self.lemmatizer.lemmatise(variant)
                if analyses and len(analyses) > 0:
                    return True
            except Exception:
                pass

        return False

    def analyze_page_xml(self, input_path, column_mode='single'):
        """
        Analyse un fichier ou dossier XML Pages.

        Args:
            input_path (str): Chemin vers fichier XML ou dossier
            column_mode (str): 'single' ou 'dual'

        Returns:
            dict: Statistiques et résultats de l'analyse
        """
        print(f"📄 Extraction du texte depuis XML Pages...")

        parser = PageXMLParser(column_mode=column_mode)

        # Déterminer si c'est un fichier ou dossier
        if os.path.isfile(input_path):
            lines, metadata = parser.parse_file(input_path)
            print(f"  ✅ 1 fichier traité, {len(lines)} lignes extraites")
        elif os.path.isdir(input_path):
            text, metadata_list = parser.parse_folder(input_path)
            lines = text.split('\n')
            print(f"  ✅ {len(metadata_list)} fichiers traités, {len(lines)} lignes extraites")
        else:
            raise ValueError(f"Chemin invalide : {input_path}")

        # Analyser le texte extrait
        return self._analyze_lines(lines, source=input_path)

    def analyze_text_file(self, input_file):
        """
        Analyse un fichier texte brut complet.

        Args:
            input_file (str): Chemin vers le fichier texte

        Returns:
            dict: Statistiques et résultats de l'analyse
        """
        print(f"📄 Analyse du fichier texte : {input_file}")

        with open(input_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        return self._analyze_lines(lines, source=input_file)

    def _analyze_lines(self, lines, source):
        """
        Analyse une liste de lignes de texte.

        Args:
            lines (list): Liste des lignes à analyser
            source (str): Source du texte (pour logs)

        Returns:
            dict: Statistiques et résultats de l'analyse
        """
        print(f"🔍 Analyse en cours...")
        print(f"   Source : {source}")

        results = []
        word_counter = Counter()
        score_distribution = {'black': 0, 'orange': 0, 'red': 0}

        total_words = 0

        for line_num, line in enumerate(lines, 1):
            # Extraire les mots de la ligne
            words = re.findall(r'\b[a-zA-ZÀ-ÿ]+\b', line)

            for i, word in enumerate(words):
                if len(word) < 2:
                    continue

                total_words += 1

                # Contexte = mots environnants (±3 mots)
                context_start = max(0, i - 3)
                context_end = min(len(words), i + 4)
                context = words[context_start:i] + words[i+1:context_end]

                # Analyser le mot
                analysis = self.analyze_word(word, context)

                results.append({
                    'line': line_num,
                    'word': word,
                    'analysis': analysis
                })

                word_counter[word.lower()] += 1
                score_distribution[analysis['color_code']] += 1

                # Afficher progression tous les 1000 mots
                if total_words % 1000 == 0:
                    print(f"  ⏳ {total_words} mots analysés...")

        print(f"✅ Analyse terminée : {total_words} mots traités")
        print(f"\n📊 Distribution des scores :")
        print(f"  ✅ Noir (bons mots)      : {score_distribution['black']} ({score_distribution['black']*100//total_words}%)")
        print(f"  ⚠️  Orange (douteux)      : {score_distribution['orange']} ({score_distribution['orange']*100//total_words}%)")
        print(f"  ❌ Rouge (erreurs prob.) : {score_distribution['red']} ({score_distribution['red']*100//total_words}%)")

        return {
            'results': results,
            'statistics': {
                'total_words': total_words,
                'distribution': score_distribution,
                'word_frequency': word_counter
            }
        }

    def generate_docx(self, input_source, output_docx, analysis_results):
        """
        Génère un document Word avec colorisation à 3 niveaux.

        Args:
            input_source (str or list): Fichier texte source OU liste de lignes
            output_docx (str): Fichier DOCX de sortie
            analysis_results (dict): Résultats de l'analyse
        """
        print(f"\n📝 Génération du document Word...")

        # Lire le fichier ou utiliser les lignes fournies
        if isinstance(input_source, str):
            with open(input_source, 'r', encoding='utf-8') as f:
                original_lines = f.readlines()
        elif isinstance(input_source, list):
            # S'assurer que chaque ligne se termine par \n
            original_lines = [line if line.endswith('\n') else line + '\n'
                             for line in input_source]
        else:
            raise ValueError("input_source doit être un chemin de fichier ou une liste de lignes")

        # Créer un index des analyses par ligne et mot
        analysis_index = {}
        for item in analysis_results['results']:
            line_num = item['line']
            word = item['word'].lower()
            if line_num not in analysis_index:
                analysis_index[line_num] = {}
            analysis_index[line_num][word] = item['analysis']

        # Créer le document
        doc = docx.Document()

        # Configuration de la page
        section = doc.sections[0]
        section.page_width = Cm(40)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        # Ajouter un titre
        title = doc.add_paragraph()
        title_run = title.add_run("Analyse de texte latin médiéval")
        title_run.bold = True
        title_run.font.size = Pt(14)

        # Légende
        legend = doc.add_paragraph()
        legend.add_run("Légende : ").bold = True

        black_run = legend.add_run("Noir = OK (score ≥75) ")
        black_run.font.color.rgb = RGBColor(0, 0, 0)

        orange_run = legend.add_run("Orange = À vérifier (score 40-74) ")
        orange_run.font.color.rgb = RGBColor(255, 140, 0)

        red_run = legend.add_run("Rouge = Erreur probable (score <40)")
        red_run.font.color.rgb = RGBColor(255, 0, 0)

        doc.add_paragraph("_" * 80)

        # Traiter chaque ligne
        for line_num, line in enumerate(original_lines, 1):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # Découper en mots et espaces
            tokens = re.findall(r'(\S+|\s+)', line)

            for token in tokens:
                if token.isspace():
                    paragraph.add_run(token)
                else:
                    # Extraire le mot pur (sans ponctuation)
                    clean_token = ''.join(c for c in token if c.isalpha()).lower()

                    # Récupérer l'analyse
                    color = RGBColor(0, 0, 0)  # Noir par défaut
                    if line_num in analysis_index and clean_token in analysis_index[line_num]:
                        analysis = analysis_index[line_num][clean_token]
                        if analysis['color_code'] == 'red':
                            color = RGBColor(255, 0, 0)
                        elif analysis['color_code'] == 'orange':
                            color = RGBColor(255, 140, 0)

                    run = paragraph.add_run(token)
                    run.font.color.rgb = color
                    run.font.size = Pt(10)

        # Sauvegarder
        doc.save(output_docx)
        print(f"✅ Document créé : {output_docx}")


def main_xml_pages():
    """
    Exemple d'utilisation avec des fichiers XML Pages.

    À adapter selon vos fichiers !
    """
    print("=" * 70)
    print("  ANALYSEUR DE TEXTES LATINS - MODE XML PAGES")
    print("  PyCollatinus + Du Cange + Extraction MainZone")
    print("=" * 70)
    print()

    # ⚙️  ADAPTER CES CHEMINS À VOTRE STRUCTURE ⚙️
    xml_input = "/path/to/xml_pages_folder"  # Dossier de fichiers XML Pages
    output_docx = "/path/to/output.docx"
    column_mode = 'single'  # ou 'dual' si vos pages ont 2 colonnes
    ducange_dict = "/home/user/Data_Base/ducange_data/dictionnaire_ducange.txt"

    # Vérifier que le chemin existe
    if not os.path.exists(xml_input):
        print(f"❌ Chemin XML non trouvé : {xml_input}")
        print("⚙️  Modifiez les chemins dans le script main_xml_pages()")
        return 1

    # Initialiser l'analyseur
    analyzer = LatinAnalyzer(ducange_dict_file=ducange_dict)

    # Analyser depuis XML Pages (avec extraction MainZone automatique)
    analysis_results = analyzer.analyze_page_xml(xml_input, column_mode=column_mode)

    # Récupérer les lignes extraites pour le DOCX
    parser = PageXMLParser(column_mode=column_mode)
    if os.path.isfile(xml_input):
        lines, _ = parser.parse_file(xml_input)
    else:
        text, _ = parser.parse_folder(xml_input)
        lines = text.split('\n')

    # Générer le document Word
    analyzer.generate_docx(lines, output_docx, analysis_results)

    print("\n" + "=" * 70)
    print("✅ TRAITEMENT TERMINÉ !")
    print(f"📁 Fichier généré : {output_docx}")
    print("=" * 70)

    return 0


def main():
    """Fonction principale pour fichier texte brut."""
    print("=" * 70)
    print("  ANALYSEUR DE TEXTES LATINS MÉDIÉVAUX - VERSION 2")
    print("  PyCollatinus + Du Cange + Scoring Multi-critères")
    print("=" * 70)
    print()

    # Chemins par défaut (à adapter)
    default_input = "/home/titouan/Téléchargements/Arras/resultats/synthese_arborescence.txt"
    default_output = "/home/titouan/Téléchargements/Arras_v2.docx"
    default_ducange = "/home/user/Data_Base/ducange_data/dictionnaire_ducange.txt"

    # Vérifier que les fichiers existent
    if not os.path.exists(default_input):
        print(f"❌ Fichier d'entrée non trouvé : {default_input}")
        print("⚙️  Modifiez les chemins dans le script (ligne ~420)")
        print("\n💡 Pour analyser des XML Pages, utilisez main_xml_pages() à la place")
        return 1

    # Initialiser l'analyseur
    analyzer = LatinAnalyzer(ducange_dict_file=default_ducange)

    # Analyser le texte
    analysis_results = analyzer.analyze_text_file(default_input)

    # Générer le document Word
    analyzer.generate_docx(default_input, default_output, analysis_results)

    print("\n" + "=" * 70)
    print("✅ TRAITEMENT TERMINÉ !")
    print(f"📁 Fichier généré : {default_output}")
    print("=" * 70)

    return 0


if __name__ == "__main__":
    sys.exit(main())
